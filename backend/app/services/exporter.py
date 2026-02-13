# backend/app/services/exporter.py
# PDF and DOCX report export service
# Generates formatted Lithuanian procurement analysis reports
# Related: models/schemas.py (AggregatedReport, QAEvaluation)

import logging
from pathlib import Path
import tempfile
from datetime import datetime

from app.models.schemas import (
    AggregatedReport,
    ConfidenceNote,
    EstimatedValue,
    QAEvaluation,
)

logger = logging.getLogger(__name__)

# Lithuanian month names (genitive case for dates)
_LT_MONTHS_GENITIVE = {
    1: "sausio",
    2: "vasario",
    3: "kovo",
    4: "balandžio",
    5: "gegužės",
    6: "birželio",
    7: "liepos",
    8: "rugpjūčio",
    9: "rugsėjo",
    10: "spalio",
    11: "lapkričio",
    12: "gruodžio",
}

_NOT_SPECIFIED = "Nenurodyta"


# ── Helper functions ───────────────────────────────────────────────────────────


def _format_value(value: EstimatedValue | None) -> str:
    """Format estimated value: '125,000.00 EUR (su PVM)' or 'Nenurodyta'."""
    if value is None or value.amount is None:
        return _NOT_SPECIFIED
    formatted = f"{value.amount:,.2f} {value.currency}"
    if value.vat_included is True:
        formatted += " (su PVM)"
    elif value.vat_included is False:
        formatted += " (be PVM)"
    if value.vat_amount is not None:
        formatted += f", PVM: {value.vat_amount:,.2f} {value.currency}"
    return formatted


def _format_date(date_str: str | None) -> str:
    """Format ISO date to Lithuanian: '2026-03-15' → '2026 m. kovo 15 d.' or 'Nenurodyta'."""
    if not date_str:
        return _NOT_SPECIFIED
    try:
        dt = datetime.fromisoformat(date_str)
        month_name = _LT_MONTHS_GENITIVE.get(dt.month, str(dt.month))
        return f"{dt.year} m. {month_name} {dt.day} d."
    except (ValueError, TypeError):
        return date_str  # Return raw string if can't parse


def _qa_color(score: float) -> str:
    """Return color name for QA score."""
    if score > 0.8:
        return "green"
    elif score > 0.5:
        return "orange"
    else:
        return "red"


def _severity_color(severity: str) -> tuple:
    """Return RGB tuple for confidence note severity."""
    mapping = {
        "info": (0, 0, 180),       # Blue
        "warning": (200, 150, 0),   # Orange/yellow
        "conflict": (200, 0, 0),    # Red
    }
    return mapping.get(severity.lower(), (0, 0, 0))


def _or_na(value: str | None) -> str:
    """Return value or 'Nenurodyta' if None/empty."""
    return value if value else _NOT_SPECIFIED


def _parse_confidence_notes(notes: list) -> list[ConfidenceNote]:
    """Parse confidence_notes which may be strings or ConfidenceNote objects."""
    result = []
    for note in notes:
        if isinstance(note, ConfidenceNote):
            result.append(note)
        elif isinstance(note, str):
            result.append(ConfidenceNote(note=note, severity="info"))
        elif isinstance(note, dict):
            result.append(ConfidenceNote(**note))
        else:
            result.append(ConfidenceNote(note=str(note), severity="info"))
    return result


# ── PDF Export ─────────────────────────────────────────────────────────────────


async def export_pdf(
    report: AggregatedReport,
    qa: QAEvaluation,
    model_used: str = "",
) -> Path:
    """
    Generate PDF report using reportlab.
    Returns path to generated PDF file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor, Color
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # Create temp file
    tmp = tempfile.NamedTemporaryFile(
        suffix=".pdf", prefix="procurement_report_", delete=False
    )
    tmp.close()
    pdf_path = Path(tmp.name)

    logger.info("Generating PDF report: %s", pdf_path)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=25 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(
        ParagraphStyle(
            "TitleLT",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "SubtitleLT",
            parent=styles["Normal"],
            fontSize=10,
            textColor=HexColor("#666666"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            "Heading2LT",
            parent=styles["Heading2"],
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "BulletLT",
            parent=styles["Normal"],
            leftIndent=20,
            bulletIndent=10,
            spaceBefore=2,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            "FooterLT",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#999999"),
            alignment=TA_CENTER,
        )
    )

    elements = []
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Title
    elements.append(Paragraph("Viešojo pirkimo analizė", styles["TitleLT"]))
    subtitle = f"Sugeneruota: {now_str}"
    if model_used:
        subtitle += f" | Modelis: {model_used}"
    elements.append(Paragraph(subtitle, styles["SubtitleLT"]))
    elements.append(Spacer(1, 6 * mm))

    # ── 1. Projekto santrauka
    elements.append(Paragraph("1. Projekto santrauka", styles["Heading2LT"]))
    elements.append(Paragraph(_or_na(report.project_summary), styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 2. Perkančioji organizacija
    elements.append(
        Paragraph("2. Perkančioji organizacija", styles["Heading2LT"])
    )
    org = report.procuring_organization
    if org:
        elements.append(
            Paragraph(f"<b>Pavadinimas:</b> {_or_na(org.name)}", styles["Normal"])
        )
        elements.append(
            Paragraph(f"<b>Kodas:</b> {_or_na(org.code)}", styles["Normal"])
        )
        elements.append(
            Paragraph(f"<b>Kontaktai:</b> {_or_na(org.contact)}", styles["Normal"])
        )
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 3. Pirkimo būdas
    if report.procurement_type:
        elements.append(Paragraph("3. Pirkimo būdas", styles["Heading2LT"]))
        elements.append(Paragraph(report.procurement_type, styles["Normal"]))
        elements.append(Spacer(1, 3 * mm))

    # ── 4. Pirkimo vertė
    elements.append(Paragraph("4. Pirkimo vertė", styles["Heading2LT"]))
    elements.append(
        Paragraph(_format_value(report.estimated_value), styles["Normal"])
    )
    elements.append(Spacer(1, 3 * mm))

    # ── 5. Terminai
    elements.append(Paragraph("5. Terminai", styles["Heading2LT"]))
    dl = report.deadlines
    if dl:
        elements.append(
            Paragraph(
                f"<b>Pasiūlymų pateikimas:</b> {_format_date(dl.submission_deadline)}",
                styles["Normal"],
            )
        )
        elements.append(
            Paragraph(
                f"<b>Klausimų pateikimas:</b> {_format_date(dl.questions_deadline)}",
                styles["Normal"],
            )
        )
        elements.append(
            Paragraph(
                f"<b>Sutarties trukmė:</b> {_or_na(dl.contract_duration)}",
                styles["Normal"],
            )
        )
        elements.append(
            Paragraph(
                f"<b>Darbų atlikimas:</b> {_format_date(dl.execution_deadline)}",
                styles["Normal"],
            )
        )
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 6. Pagrindiniai reikalavimai
    elements.append(
        Paragraph("6. Pagrindiniai reikalavimai", styles["Heading2LT"])
    )
    if report.key_requirements:
        for req in report.key_requirements:
            elements.append(
                Paragraph(f"• {req}", styles["BulletLT"])
            )
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 7. Kvalifikacijos reikalavimai
    elements.append(
        Paragraph("7. Kvalifikacijos reikalavimai", styles["Heading2LT"])
    )
    qr = report.qualification_requirements
    if qr:
        for group_name, group_label in [
            ("financial", "Finansiniai"),
            ("technical", "Techniniai"),
            ("experience", "Patirties"),
            ("other", "Kiti"),
        ]:
            items = getattr(qr, group_name, [])
            if items:
                elements.append(
                    Paragraph(f"<b>{group_label}:</b>", styles["Normal"])
                )
                for item in items:
                    elements.append(
                        Paragraph(f"• {item}", styles["BulletLT"])
                    )
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 8. Vertinimo kriterijai (TABLE)
    elements.append(
        Paragraph("8. Vertinimo kriterijai", styles["Heading2LT"])
    )
    if report.evaluation_criteria:
        table_data = [["Kriterijus", "Svoris (%)", "Aprašymas"]]
        for ec in report.evaluation_criteria:
            weight = f"{ec.weight_percent:.1f}" if ec.weight_percent is not None else "-"
            table_data.append([ec.criterion, weight, _or_na(ec.description)])

        col_widths = [150, 60, 260]
        table = Table(table_data, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("ALIGN", (1, 0), (1, -1), "CENTER"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#F2F2F2")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        elements.append(table)
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 9. Lotai (TABLE)
    elements.append(Paragraph("9. Lotai", styles["Heading2LT"]))
    if report.lot_structure:
        lot_data = [["Nr.", "Aprašymas", "Vertė (EUR)"]]
        for lot in report.lot_structure:
            val = f"{lot.estimated_value:,.2f}" if lot.estimated_value is not None else "-"
            lot_data.append([str(lot.lot_number), lot.description, val])

        lot_table = Table(lot_data, colWidths=[40, 320, 100])
        lot_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("ALIGN", (0, 0), (0, -1), "CENTER"),
                    ("ALIGN", (2, 0), (2, -1), "RIGHT"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#F2F2F2")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(lot_table)
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 10. Specialios sąlygos
    elements.append(
        Paragraph("10. Specialios sąlygos", styles["Heading2LT"])
    )
    if report.special_conditions:
        for cond in report.special_conditions:
            elements.append(Paragraph(f"• {cond}", styles["BulletLT"]))
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 11. Apribojimai ir draudimai
    elements.append(
        Paragraph("11. Apribojimai ir draudimai", styles["Heading2LT"])
    )
    if report.restrictions_and_prohibitions:
        for r in report.restrictions_and_prohibitions:
            elements.append(Paragraph(f"• {r}", styles["BulletLT"]))
    else:
        elements.append(Paragraph(_NOT_SPECIFIED, styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 12. Pastabos ir patikimumas
    elements.append(
        Paragraph("12. Pastabos ir patikimumas", styles["Heading2LT"])
    )
    if report.confidence_notes:
        notes = _parse_confidence_notes(report.confidence_notes)
        for cn in notes:
            r, g, b = _severity_color(cn.severity)
            color_hex = f"#{r:02x}{g:02x}{b:02x}"
            elements.append(
                Paragraph(
                    f'<font color="{color_hex}">[{cn.severity.upper()}]</font> {cn.note}',
                    styles["Normal"],
                )
            )
    else:
        elements.append(Paragraph("Pastabų nėra", styles["Normal"]))
    elements.append(Spacer(1, 3 * mm))

    # ── 13. Kokybės vertinimas (QA)
    elements.append(
        Paragraph("13. Kokybės vertinimas", styles["Heading2LT"])
    )
    qa_color = _qa_color(qa.completeness_score)
    color_map = {"green": "#2E7D32", "orange": "#E65100", "red": "#C62828"}
    qa_hex = color_map.get(qa_color, "#000000")

    elements.append(
        Paragraph(
            f'<b>Užbaigtumo balas:</b> <font color="{qa_hex}">'
            f"{qa.completeness_score:.0%}</font>",
            styles["Normal"],
        )
    )

    if qa.missing_fields:
        elements.append(
            Paragraph("<b>Trūkstami laukai:</b>", styles["Normal"])
        )
        for mf in qa.missing_fields:
            elements.append(Paragraph(f"• {mf}", styles["BulletLT"]))

    if qa.conflicts:
        elements.append(
            Paragraph("<b>Prieštaravimai:</b>", styles["Normal"])
        )
        for conflict in qa.conflicts:
            elements.append(
                Paragraph(
                    f'<font color="#C62828">• {conflict}</font>',
                    styles["BulletLT"],
                )
            )

    if qa.suggestions:
        elements.append(
            Paragraph("<b>Pasiūlymai:</b>", styles["Normal"])
        )
        for sug in qa.suggestions:
            elements.append(Paragraph(f"• {sug}", styles["BulletLT"]))

    elements.append(Spacer(1, 6 * mm))

    # ── Footer
    footer_text = f"Sugeneruota: {now_str}"
    if model_used:
        footer_text += f" | Modelis: {model_used}"
    elements.append(Paragraph(footer_text, styles["FooterLT"]))

    # Build
    doc.build(elements)
    logger.info("PDF report generated: %s (%d bytes)", pdf_path, pdf_path.stat().st_size)
    return pdf_path


# ── DOCX Export ────────────────────────────────────────────────────────────────


async def export_docx(
    report: AggregatedReport,
    qa: QAEvaluation,
    model_used: str = "",
) -> Path:
    """
    Generate DOCX report using python-docx.
    Returns path to generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx", prefix="procurement_report_", delete=False
    )
    tmp.close()
    docx_path = Path(tmp.name)

    logger.info("Generating DOCX report: %s", docx_path)

    doc = Document()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Title
    title_para = doc.add_heading("Viešojo pirkimo analizė", level=0)
    subtitle = f"Sugeneruota: {now_str}"
    if model_used:
        subtitle += f" | Modelis: {model_used}"
    sub_para = doc.add_paragraph(subtitle)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(102, 102, 102)

    # ── 1. Projekto santrauka
    doc.add_heading("1. Projekto santrauka", level=1)
    doc.add_paragraph(_or_na(report.project_summary))

    # ── 2. Perkančioji organizacija
    doc.add_heading("2. Perkančioji organizacija", level=1)
    org = report.procuring_organization
    if org:
        doc.add_paragraph(f"Pavadinimas: {_or_na(org.name)}")
        doc.add_paragraph(f"Kodas: {_or_na(org.code)}")
        doc.add_paragraph(f"Kontaktai: {_or_na(org.contact)}")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 3. Pirkimo būdas
    if report.procurement_type:
        doc.add_heading("3. Pirkimo būdas", level=1)
        doc.add_paragraph(report.procurement_type)

    # ── 4. Pirkimo vertė
    doc.add_heading("4. Pirkimo vertė", level=1)
    doc.add_paragraph(_format_value(report.estimated_value))

    # ── 5. Terminai
    doc.add_heading("5. Terminai", level=1)
    dl = report.deadlines
    if dl:
        doc.add_paragraph(
            f"Pasiūlymų pateikimas: {_format_date(dl.submission_deadline)}"
        )
        doc.add_paragraph(
            f"Klausimų pateikimas: {_format_date(dl.questions_deadline)}"
        )
        doc.add_paragraph(
            f"Sutarties trukmė: {_or_na(dl.contract_duration)}"
        )
        doc.add_paragraph(
            f"Darbų atlikimas: {_format_date(dl.execution_deadline)}"
        )
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 6. Pagrindiniai reikalavimai
    doc.add_heading("6. Pagrindiniai reikalavimai", level=1)
    if report.key_requirements:
        for req in report.key_requirements:
            doc.add_paragraph(req, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 7. Kvalifikacijos reikalavimai
    doc.add_heading("7. Kvalifikacijos reikalavimai", level=1)
    qr = report.qualification_requirements
    if qr:
        for group_name, group_label in [
            ("financial", "Finansiniai"),
            ("technical", "Techniniai"),
            ("experience", "Patirties"),
            ("other", "Kiti"),
        ]:
            items = getattr(qr, group_name, [])
            if items:
                p = doc.add_paragraph()
                run = p.add_run(f"{group_label}:")
                run.bold = True
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 8. Vertinimo kriterijai (TABLE)
    doc.add_heading("8. Vertinimo kriterijai", level=1)
    if report.evaluation_criteria:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Kriterijus"
        hdr[1].text = "Svoris (%)"
        hdr[2].text = "Aprašymas"

        for ec in report.evaluation_criteria:
            row = table.add_row().cells
            row[0].text = ec.criterion
            row[1].text = (
                f"{ec.weight_percent:.1f}" if ec.weight_percent is not None else "-"
            )
            row[2].text = _or_na(ec.description)
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 9. Lotai (TABLE)
    doc.add_heading("9. Lotai", level=1)
    if report.lot_structure:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Nr."
        hdr[1].text = "Aprašymas"
        hdr[2].text = "Vertė (EUR)"

        for lot in report.lot_structure:
            row = table.add_row().cells
            row[0].text = str(lot.lot_number)
            row[1].text = lot.description
            row[2].text = (
                f"{lot.estimated_value:,.2f}" if lot.estimated_value is not None else "-"
            )
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 10. Specialios sąlygos
    doc.add_heading("10. Specialios sąlygos", level=1)
    if report.special_conditions:
        for cond in report.special_conditions:
            doc.add_paragraph(cond, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 11. Apribojimai ir draudimai
    doc.add_heading("11. Apribojimai ir draudimai", level=1)
    if report.restrictions_and_prohibitions:
        for r in report.restrictions_and_prohibitions:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 12. Pastabos ir patikimumas
    doc.add_heading("12. Pastabos ir patikimumas", level=1)
    if report.confidence_notes:
        notes = _parse_confidence_notes(report.confidence_notes)
        for cn in notes:
            p = doc.add_paragraph()
            severity_run = p.add_run(f"[{cn.severity.upper()}] ")
            r, g, b = _severity_color(cn.severity)
            severity_run.font.color.rgb = RGBColor(r, g, b)
            severity_run.bold = True
            p.add_run(cn.note)
    else:
        doc.add_paragraph("Pastabų nėra")

    # ── 13. Kokybės vertinimas
    doc.add_heading("13. Kokybės vertinimas", level=1)

    # QA Score with color
    qa_p = doc.add_paragraph()
    qa_p.add_run("Užbaigtumo balas: ").bold = True
    score_run = qa_p.add_run(f"{qa.completeness_score:.0%}")
    score_run.bold = True
    qa_color = _qa_color(qa.completeness_score)
    color_map = {
        "green": RGBColor(0x2E, 0x7D, 0x32),
        "orange": RGBColor(0xE6, 0x51, 0x00),
        "red": RGBColor(0xC6, 0x28, 0x28),
    }
    score_run.font.color.rgb = color_map.get(qa_color, RGBColor(0, 0, 0))

    if qa.missing_fields:
        p = doc.add_paragraph()
        p.add_run("Trūkstami laukai:").bold = True
        for mf in qa.missing_fields:
            doc.add_paragraph(mf, style="List Bullet")

    if qa.conflicts:
        p = doc.add_paragraph()
        p.add_run("Prieštaravimai:").bold = True
        for conflict in qa.conflicts:
            cp = doc.add_paragraph(style="List Bullet")
            run = cp.add_run(conflict)
            run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    if qa.suggestions:
        p = doc.add_paragraph()
        p.add_run("Pasiūlymai:").bold = True
        for sug in qa.suggestions:
            doc.add_paragraph(sug, style="List Bullet")

    # ── Footer
    doc.add_paragraph()  # spacer
    footer_text = f"Sugeneruota: {now_str}"
    if model_used:
        footer_text += f" | Modelis: {model_used}"
    footer_para = doc.add_paragraph(footer_text)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(153, 153, 153)

    doc.save(str(docx_path))
    logger.info("DOCX report generated: %s (%d bytes)", docx_path, docx_path.stat().st_size)
    return docx_path

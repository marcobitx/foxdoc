# backend/tests/test_parser.py
# Tests for the document parsing service (services/parser.py)
# Covers: Docling conversion, classification heuristics, page estimation, error handling

import asyncio
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument

from app.models.schemas import DocumentType
from app.services.parser import (
    ParsedDocument,
    _estimate_pages,
    classify_document,
    parse_all,
    parse_document,
)


# ── Fixtures ──────────────────────────────────────────────────────────────────


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a simple DOCX file for testing."""
    doc = DocxDocument()
    doc.add_heading("Techninė specifikacija", level=1)
    doc.add_paragraph(
        "Ši techninė specifikacija aprašo reikalavimus pirkimui. "
        "Perkančioji organizacija numato įsigyti programinę įrangą."
    )
    doc.add_heading("Reikalavimai", level=2)
    for i in range(1, 6):
        doc.add_paragraph(f"Reikalavimas {i}: Lorem ipsum dolor sit amet, " * 5)
    path = tmp_path / "techninė_specifikacija.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def minimal_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with just one line."""
    doc = DocxDocument()
    doc.add_paragraph("Hello world")
    path = tmp_path / "minimal.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def corrupt_file(tmp_path: Path) -> Path:
    """Create a corrupt file that Docling can't parse."""
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"this is not a valid PDF file at all\x00\xff\xfe")
    return path


# ── Document parsing tests ────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_parse_docx(sample_docx: Path):
    """Test parsing a real DOCX file with Docling."""
    result = await parse_document(sample_docx, "techninė_specifikacija.docx")

    assert isinstance(result, ParsedDocument)
    assert result.filename == "techninė_specifikacija.docx"
    assert len(result.content) > 0
    assert "[ERROR]" not in result.content
    assert result.file_size_bytes > 0
    assert result.page_count >= 1
    assert result.token_estimate > 0
    assert result.doc_type == DocumentType.TECHNICAL_SPEC


@pytest.mark.asyncio
async def test_parse_minimal_docx(minimal_docx: Path):
    """Test parsing a minimal DOCX file."""
    result = await parse_document(minimal_docx, "minimal.docx")

    assert isinstance(result, ParsedDocument)
    assert "Hello world" in result.content or "Hello" in result.content
    assert result.page_count >= 1
    assert result.doc_type == DocumentType.OTHER


@pytest.mark.asyncio
async def test_parse_corrupt_file(corrupt_file: Path):
    """Test that corrupt files return error content instead of crashing."""
    result = await parse_document(corrupt_file, "corrupt.pdf")

    assert isinstance(result, ParsedDocument)
    assert "[ERROR]" in result.content
    assert result.filename == "corrupt.pdf"
    assert result.file_size_bytes > 0
    assert result.page_count == 0


@pytest.mark.asyncio
async def test_parse_nonexistent_file():
    """Test that a missing file returns an error gracefully."""
    result = await parse_document(
        Path("/nonexistent/file.docx"), "missing.docx"
    )

    assert isinstance(result, ParsedDocument)
    assert "[ERROR]" in result.content


# ── parse_all tests ──────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_parse_all(sample_docx: Path, minimal_docx: Path):
    """Test parsing multiple files."""
    file_paths = [
        (sample_docx, "techninė_specifikacija.docx"),
        (minimal_docx, "minimal.docx"),
    ]

    results = await parse_all(file_paths)

    assert len(results) == 2
    assert results[0].filename == "techninė_specifikacija.docx"
    assert results[1].filename == "minimal.docx"


@pytest.mark.asyncio
async def test_parse_all_with_callback(sample_docx: Path):
    """Test that the on_parsed callback is invoked."""
    callback_results: list[ParsedDocument] = []

    results = await parse_all(
        [(sample_docx, "test.docx")],
        on_parsed=lambda doc: callback_results.append(doc),
    )

    assert len(results) == 1
    assert len(callback_results) == 1
    assert callback_results[0].filename == "test.docx"


@pytest.mark.asyncio
async def test_parse_all_empty():
    """Test parsing an empty list."""
    results = await parse_all([])
    assert results == []


# ── Classification tests ─────────────────────────────────────────────────────


class TestClassifyDocument:
    """Test document type classification from filename and content."""

    def test_technical_spec_from_filename(self):
        assert (
            classify_document("Techninė_specifikacija.pdf", "")
            == DocumentType.TECHNICAL_SPEC
        )

    def test_technical_spec_technin(self):
        assert (
            classify_document("techninis_aprasymas.docx", "")
            == DocumentType.TECHNICAL_SPEC
        )

    def test_contract_from_filename(self):
        assert (
            classify_document("sutartis_nr_123.pdf", "")
            == DocumentType.CONTRACT
        )

    def test_contract_case_insensitive(self):
        assert (
            classify_document("SUTARTIS.PDF", "")
            == DocumentType.CONTRACT
        )

    def test_invitation_kvietimas(self):
        assert (
            classify_document("kvietimas_dalyvauti.pdf", "")
            == DocumentType.INVITATION
        )

    def test_invitation_skelbimas(self):
        assert (
            classify_document("skelbimas_apie_pirkima.pdf", "")
            == DocumentType.INVITATION
        )

    def test_qualification_from_filename(self):
        assert (
            classify_document("kvalifikacijos_reikalavimai.pdf", "")
            == DocumentType.QUALIFICATION
        )

    def test_evaluation_vertinimas(self):
        assert (
            classify_document("vertinimo_kriterijai.xlsx", "")
            == DocumentType.EVALUATION
        )

    def test_evaluation_kriteriju(self):
        assert (
            classify_document("kriteriju_lentele.pdf", "")
            == DocumentType.EVALUATION
        )

    def test_annex_priedas(self):
        assert (
            classify_document("priedas_nr_1.docx", "")
            == DocumentType.ANNEX
        )

    def test_annex_forma(self):
        assert (
            classify_document("forma_pasiulymo.docx", "")
            == DocumentType.ANNEX
        )

    def test_annex_sablonas(self):
        assert (
            classify_document("šablonas_dokumentui.docx", "")
            == DocumentType.ANNEX
        )

    def test_other_default(self):
        assert (
            classify_document("random_document.pdf", "")
            == DocumentType.OTHER
        )

    def test_fallback_to_content(self):
        """When filename doesn't match, content preview should be checked."""
        assert (
            classify_document(
                "document_001.pdf",
                "Ši techninė specifikacija aprašo reikalavimus...",
            )
            == DocumentType.TECHNICAL_SPEC
        )

    def test_content_contract(self):
        assert (
            classify_document(
                "file.pdf",
                "Sutarties sąlygos ir priedai",
            )
            == DocumentType.CONTRACT
        )

    def test_filename_takes_priority(self):
        """Filename classification should win over content classification."""
        assert (
            classify_document(
                "sutartis.pdf",
                "techninė specifikacija text here",
            )
            == DocumentType.CONTRACT
        )


# ── Page estimation tests ────────────────────────────────────────────────────


class TestEstimatePages:
    def test_empty_content(self):
        assert _estimate_pages("", ".pdf") == 0

    def test_short_content_one_page(self):
        assert _estimate_pages("Short text", ".pdf") == 1

    def test_long_content_multiple_pages(self):
        content = "x" * 9000  # Should be ~3 pages
        assert _estimate_pages(content, ".pdf") == 3

    def test_xlsx_single_sheet(self):
        content = "Some table data without headers"
        assert _estimate_pages(content, ".xlsx") == 1

    def test_xlsx_multiple_sheets(self):
        content = "## Sheet 1\nData\n## Sheet 2\nMore data\n## Sheet 3\nEven more"
        assert _estimate_pages(content, ".xlsx") == 3

    def test_docx_estimation(self):
        content = "x" * 6000  # ~2 pages
        assert _estimate_pages(content, ".docx") == 2

    def test_minimum_one_page_for_nonempty(self):
        assert _estimate_pages("a", ".pdf") == 1
